# This Python file uses the following encoding: utf-8

from docxtpl import DocxTemplate
from docx.shared import Inches

class Word():

    def generate_reports(self, evaluations, file_naming_convention):

        # get list of all students
        students = ()
        for evaluation in evaluations:
            for student in evaluation:
                students += (student,)

        students = sorted(set(students))

        for student_name in students:


            student_evaluations = []
            for evaluation in evaluations:
                if student_name in evaluation:
                    student_evaluations.append(evaluation)

            doc = DocxTemplate("templates/report.docx")

            context = { 
                'student_name' : student_name.decode('utf8'),
                'total_grade': evaluations[0][student_name]['total_grade'],
                'teacher_name': evaluations[0][student_name]['teacher_name'],
                'area_name': file_naming_convention(evaluations[0][student_name]['area_name']),
                'trimester_number': evaluations[0][student_name]['trimester_number'],

            }

            filename = file_naming_convention(student_name)

            for paragraph in doc.paragraphs:
                if "[image]" in paragraph.text:
                    paragraph.text = paragraph.text.strip().replace("[image]", "")

                    run = paragraph.add_run()
                    run.add_picture('documents/images-output/%s.png' % filename, width=Inches(3))


            doc.render(context)
            doc.save("documents/word-output/%s.docx" % filename)